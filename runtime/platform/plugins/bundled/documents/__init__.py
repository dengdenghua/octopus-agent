"""documents — 本地 docx 文档处理插件(独立自研)。

能力范围对标 OpenAI codex "documents" 插件(创建/编辑/提取 docx),但为
Octopus 完全原创实现:
- 仅依赖通用开源库 python-docx + lxml;
- 不包含、不复制 OpenAI 的任何代码或提示词文本;
- 无任何外部服务 / connector / OAuth 依赖,全部本地完成。

技能:
  documents.create_docx   从结构化内容创建 .docx(写操作,覆盖需 overwrite=true)
  documents.extract_text  提取 .docx 结构化文本(段落/表格/图片)
  documents.to_markdown   .docx -> Markdown
  documents.docx_info     文档元信息统计
"""

from __future__ import annotations

import contextlib
import logging
from pathlib import Path
from typing import Any

from runtime.execution.suckers.registry import Skill
from runtime.platform.plugins.plugin_base import ModulePlugin

_LOGGER = logging.getLogger(__name__)

PLUGIN_NAME = "documents"
_TRUSTED_SOURCE = "plugin://documents"

try:  # pragma: no cover - 依赖探测
    from docx import Document  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from docx.table import Table  # type: ignore[import-untyped]
    from docx.text.paragraph import Paragraph  # type: ignore[import-untyped]

    _DOCX_OK = True
except Exception:  # pragma: no cover
    _DOCX_OK = False

# 提取文本上限,防止超大文档撑爆上下文
_MAX_TEXT = 200_000
# 单次返回段落/表格上限
_MAX_PARAS = 500
_MAX_TABLE_ROWS = 50


def _require_docx() -> dict[str, Any] | None:
    """python-docx 不可用时返回错误 dict。"""
    if not _DOCX_OK:
        return {
            "ok": False,
            "error": "python-docx 未安装,无法处理 docx(安装: pip install python-docx)",
        }
    return None


def _resolve_path(path: Any) -> tuple[Path | None, dict[str, Any] | None]:
    try:
        p = Path(str(path)).expanduser().resolve()
    except Exception as exc:
        return None, {"ok": False, "error": f"无效路径: {exc}"}
    if p.suffix.lower() != ".docx":
        return None, {"ok": False, "error": "仅支持 .docx 文件"}
    return p, None


def _set_cjk_font(doc: Any) -> None:
    """默认正文支持中文字体(失败静默)。"""
    with contextlib.suppress(Exception):
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "宋体")


def _heading_level(style_name: str) -> int | None:
    """'Heading 1' -> 1, 其他 -> None。"""
    if style_name.startswith("Heading"):
        with contextlib.suppress(ValueError):
            return int(style_name.split()[-1])
        return 1
    return None


def _iter_body_blocks(doc: Any):
    """按 .docx body 的原始顺序混合产出 ('paragraph', Paragraph) / ('table', Table)。

    doc.paragraphs / doc.tables 是分开的视图,会丢失段落与表格的穿插顺序;
    转换 Markdown 等需要保序的场景必须用本迭代器。
    """
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield "paragraph", Paragraph(child, doc)
        elif tag == "tbl":
            yield "table", Table(child, doc)


class DocumentsPlugin(ModulePlugin):
    name = PLUGIN_NAME
    display_name = "Documents"
    version = "0.1.0"
    description = (
        "本地 docx 文档处理插件(创建 / 提取 / 转 Markdown / 元信息)。"
        "独立自研,能力对标 OpenAI codex documents 插件,但不包含其代码,"
        "无外部服务依赖。"
    )
    author = "Octopus"

    # ── 技能注册 ────────────────────────────────────────────────
    def register_skills(self) -> None:
        if self.ctx is None:
            return
        skills = [
            Skill(
                name="documents.create_docx",
                description=(
                    "从结构化内容创建 .docx 文档(写操作)。参数:path 必填(输出 .docx 路径),"
                    "sections 必填(内容数组:heading/paragraph/list/table 四种元素),title 可选"
                    "(大标题),overwrite 可选(目标文件已存在时必须显式传 true 才覆盖,默认拒绝)。"
                    "适用于'帮我生成一份 Word 文档'、'把这段大纲做成 docx'。"
                ),
                summary="创建 .docx 文档(path+sections 必填)",
                affinity=["documents", "docx", "word", "create"],
                cost_profile="low",
                trusted_source=_TRUSTED_SOURCE,
                handler=self._create_docx,
            ),
            Skill(
                name="documents.extract_text",
                description=(
                    "提取 .docx 文档的结构化文本:段落(带 Word 样式名与标题层级)、表格数据、"
                    "图片数量。参数:path 必填(.docx 路径)。适用于'这个 Word 里写了啥'、"
                    "'把文档内容读出来'。"
                ),
                summary="提取 .docx 结构化文本(path 必填)",
                affinity=["documents", "docx", "word", "extract", "read"],
                cost_profile="low",
                trusted_source=_TRUSTED_SOURCE,
                handler=self._extract_text,
            ),
            Skill(
                name="documents.to_markdown",
                description=(
                    "把 .docx 文档转换为 Markdown 文本:Heading 标题转 #,列表转 -/1.,"
                    "表格转 Markdown 表格。参数:path 必填(.docx 路径)。适用于'把这个 Word "
                    "转成 Markdown'、'导出为 md'。"
                ),
                summary=".docx 转 Markdown(path 必填)",
                affinity=["documents", "docx", "markdown", "convert"],
                cost_profile="low",
                trusted_source=_TRUSTED_SOURCE,
                handler=self._to_markdown,
            ),
            Skill(
                name="documents.docx_info",
                description=(
                    "查看 .docx 文档的元信息:段落数、表格数(含行列)、图片数、节数、核心样式、"
                    "文件大小。参数:path 必填(.docx 路径)。适用于'这个文档多大'、'文档结构概览'。"
                ),
                summary="docx 元信息统计(path 必填)",
                affinity=["documents", "docx", "word", "info"],
                cost_profile="low",
                trusted_source=_TRUSTED_SOURCE,
                handler=self._docx_info,
            ),
        ]
        for skill in skills:
            with contextlib.suppress(Exception):
                self.ctx.register_skill(skill)

    # ── 工具实现(python-docx,本地) ─────────────────────────────
    def _create_docx(self, **kwargs: Any) -> dict[str, Any]:
        err = _require_docx()
        if err:
            return err
        p, perr = _resolve_path(kwargs.get("path"))
        if perr:
            return perr
        if not kwargs.get("overwrite") and p.exists():
            return {
                "ok": False,
                "error": f"目标文件已存在: {p},写操作需显式传 overwrite=true 才会覆盖",
                "preview": {
                    "path": str(p),
                    "exists": True,
                    "num_sections": len(kwargs.get("sections") or []),
                },
            }
        title = str(kwargs.get("title") or "")
        sections = kwargs.get("sections") or []

        doc = Document()
        _set_cjk_font(doc)
        if title:
            doc.add_heading(title, level=0)

        num_paras = 0
        num_headings = 0
        num_tables = 0
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            kind = sec.get("type", "paragraph")
            try:
                if kind == "heading":
                    level = int(sec.get("level") or 1)
                    level = min(max(level, 1), 6)
                    doc.add_heading(str(sec.get("text") or ""), level=level)
                    num_headings += 1
                elif kind == "list":
                    ordered = bool(sec.get("ordered"))
                    style = "List Number" if ordered else "List Bullet"
                    for item in sec.get("items") or []:
                        doc.add_paragraph(str(item), style=style)
                        num_paras += 1
                elif kind == "table":
                    headers = [str(h) for h in (sec.get("headers") or [])]
                    rows = [[str(c) for c in r] for r in (sec.get("rows") or [])]
                    ncols = max(len(headers), max((len(r) for r in rows), default=0))
                    ncols = max(ncols, 1)
                    table = doc.add_table(rows=0, cols=ncols)
                    table.style = "Table Grid"
                    if headers:
                        cells = table.add_row().cells
                        for i, h in enumerate(headers[:ncols]):
                            cells[i].text = h
                    for row in rows:
                        cells = table.add_row().cells
                        for i in range(ncols):
                            cells[i].text = row[i] if i < len(row) else ""
                    num_tables += 1
                else:  # paragraph
                    doc.add_paragraph(str(sec.get("text") or ""))
                    num_paras += 1
            except Exception as exc:  # noqa: BLE001
                _LOGGER.warning("create_docx 处理块失败(type=%s): %s", kind, exc)

        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(p))
        except Exception as exc:  # noqa: BLE001
            return {"ok": False, "error": f"保存 docx 失败: {exc}"}
        return {
            "ok": True,
            "path": str(p),
            "title": title,
            "num_paragraphs": num_paras,
            "num_headings": num_headings,
            "num_tables": num_tables,
        }

    def _extract_text(self, **kwargs: Any) -> dict[str, Any]:
        err = _require_docx()
        if err:
            return err
        p, perr = _resolve_path(kwargs.get("path"))
        if perr:
            return perr
        if not p.exists():
            return {"ok": False, "error": f"文件不存在: {p}"}
        try:
            doc = Document(str(p))
        except Exception as exc:  # noqa: BLE001
            return {"ok": False, "error": f"无法打开 docx: {exc}"}

        paragraphs: list[dict[str, Any]] = []
        total_chars = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            total_chars += len(text)
            style = para.style.name if para.style is not None else "Normal"
            paragraphs.append({"style": style, "level": _heading_level(style), "text": text})

        tables: list[dict[str, Any]] = []
        for tbl in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            tables.append(
                {
                    "rows": len(rows),
                    "cols": len(rows[0]) if rows else 0,
                    "data": rows[:_MAX_TABLE_ROWS],
                }
            )

        num_images = len(doc.inline_shapes)
        truncated = total_chars > _MAX_TEXT
        return {
            "ok": True,
            "path": str(p),
            "num_paragraphs": len(paragraphs),
            "num_tables": len(tables),
            "num_images": num_images,
            "total_chars": total_chars,
            "paragraphs": paragraphs[:_MAX_PARAS],
            "tables": tables,
            "truncated": truncated,
        }

    def _to_markdown(self, **kwargs: Any) -> dict[str, Any]:
        err = _require_docx()
        if err:
            return err
        p, perr = _resolve_path(kwargs.get("path"))
        if perr:
            return perr
        if not p.exists():
            return {"ok": False, "error": f"文件不存在: {p}"}
        try:
            doc = Document(str(p))
        except Exception as exc:  # noqa: BLE001
            return {"ok": False, "error": f"无法打开 docx: {exc}"}

        lines: list[str] = []
        list_counters: dict[str, int] = {}

        def flush_list() -> None:
            nonlocal lines
            if lines and lines[-1].startswith(("- ", "1. ", "2. ", "3. ", "4. ", "5. ")):
                lines.append("")

        for kind, block in _iter_body_blocks(doc):
            if kind == "paragraph":
                text = block.text.strip()
                if not text:
                    continue
                style = block.style.name if block.style is not None else "Normal"
                level = _heading_level(style)
                if level is not None:
                    flush_list()
                    lines.append("#" * min(level, 6) + " " + text)
                elif style.startswith("List Bullet"):
                    lines.append("- " + text)
                elif style.startswith("List Number"):
                    list_counters[style] = list_counters.get(style, 0) + 1
                    lines.append(f"{list_counters[style]}. " + text)
                else:
                    flush_list()
                    lines.append(text)
            else:  # table,按原位置输出
                flush_list()
                if lines:
                    lines.append("")
                rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                if not rows:
                    continue
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("|" + "|".join("---" for _ in rows[0]) + "|")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")

        md = "\n".join(lines).strip() + "\n"
        truncated = len(md) > _MAX_TEXT
        return {
            "ok": True,
            "path": str(p),
            "markdown": md[:_MAX_TEXT],
            "num_chars": len(md),
            "truncated": truncated,
        }

    def _docx_info(self, **kwargs: Any) -> dict[str, Any]:
        err = _require_docx()
        if err:
            return err
        p, perr = _resolve_path(kwargs.get("path"))
        if perr:
            return perr
        if not p.exists():
            return {"ok": False, "error": f"文件不存在: {p}"}
        try:
            doc = Document(str(p))
        except Exception as exc:  # noqa: BLE001
            return {"ok": False, "error": f"无法打开 docx: {exc}"}

        paragraphs = [x for x in doc.paragraphs if x.text.strip()]
        headings = [x for x in paragraphs if _heading_level(x.style.name) is not None]
        tables = [{"rows": len(t.rows), "cols": len(t.columns)} for t in doc.tables]
        styles = sorted({x.style.name for x in paragraphs})[:30]
        return {
            "ok": True,
            "path": str(p),
            "num_paragraphs": len(paragraphs),
            "num_headings": len(headings),
            "num_tables": len(tables),
            "tables": tables,
            "num_images": len(doc.inline_shapes),
            "num_sections": len(doc.sections),
            "core_styles": styles,
            "size_bytes": p.stat().st_size,
        }


__all__ = ["DocumentsPlugin"]
